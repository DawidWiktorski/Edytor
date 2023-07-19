import os
import re
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from docx import Document

def open_file_dialog():
    root = tk.Tk()
    root.withdraw()
    filetypes = (("Word documents", "*.docx"), ("All files", "*.*"))
    file_path = filedialog.askopenfilename(filetypes=filetypes)
    return file_path

def find_and_replace(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:

        # Usuń spacje na początku i końcu akapitu
        paragraph.text = re.sub(r'^\s+|\s+$', '', paragraph.text)

        # Usuń kropki, które są bezpośrednio przed innym znakiem interpunkcyjnym (oprócz kropki)
        paragraph.text = re.sub(r'\.(?=[!,;:?])', '', paragraph.text)

        # Usuń kropki, które są bezpośrednio za innym znakiem interpunkcyjnym (oprócz kropki)
        paragraph.text = re.sub(r'(?=[!,;:?]\.)', '', paragraph.text)

        # Usuń spacje przed znakami interpunkcyjnymi
        matches1 = re.findall(r'\s([?.!,:;])', paragraph.text)
        for match in matches1:
            paragraph.text = re.sub(r'\s([?.!,:;])', r'\1', paragraph.text)

        # Zastąp twarde spacje zwykłymi spacjami
        matches2 = re.findall('\u00A0', paragraph.text)
        for match in matches2:
            paragraph.text = paragraph.text.replace(match, ' ')

        # Skróć ciągi znaków interpunkcyjnych do jednego powtórzenia
        matches3 = re.findall(r'([^\w\s\.])\1{1,}', paragraph.text)
        for match in matches3:
            new_text = re.sub(r'([^\w\s\.])\1{1,}', r'\1', paragraph.text)
            paragraph.text = new_text

        # Dodaj spację po znaku interpunkcyjnym tylko wtedy, gdy za nim jest coś innego niż spacja lub znak interpunkcyjny
        matches17 = re.findall(r'[?.!,:;](?=[^\s?.!,:;])', paragraph.text)
        for match in matches17:
            paragraph.text = re.sub(r'([?.!,:;])(?=[^\s?.!,:;])', r'\1 ', paragraph.text)

        # Zamień trzy kropki na znak wielokropka
        matches4 = re.findall(r'\.\.\.', paragraph.text)
        for match in matches4:
            new_text = match.replace("...", "…")
            paragraph.text = paragraph.text.replace(match, new_text)

        # Usuń kropkę po znaku wielokropka
        matches5 = re.findall(r'…\.', paragraph.text)
        for match in matches5:
            new_text = match.replace("….", "…")
            paragraph.text = paragraph.text.replace(match, new_text)

        # Skróć ciągi liter do trzech powtórzeń
        matches6 = re.findall(r'([a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ])\1{3,}', paragraph.text)
        for match in matches6:
            new_text = re.sub(r'([a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ])\1{3,}', r'\1\1\1', paragraph.text, flags=re.IGNORECASE)
            paragraph.text = new_text

        # Zamień podwójne dywizy na jedną półpauzę
        matches7 = re.findall(r'--', paragraph.text)
        for match in matches7:
            new_text = '–'
            paragraph.text = paragraph.text.replace(match, new_text)

        # Zamień dywizy i pauzy na półpauzy
        matches8 = re.findall(r'[-—]', paragraph.text)
        for match in matches8:
            new_text = "–"
            paragraph.text = paragraph.text.replace(match, new_text)

        # Zamień półpauzy na dywiz, jeśli otoczone literami lub cyframi
        matches9 = re.findall(r'(?<=[^\s])–(?=[^\s])', paragraph.text)
        for match in matches9:
            new_text = match.replace('–', '-')
            paragraph.text = paragraph.text.replace(match, new_text)

        # Dodaj spację po półpauzie
        matches10 = re.findall(r'–(?=[^\s])', paragraph.text)
        for match in matches10:
            new_text = "– "
            paragraph.text = paragraph.text.replace(match, new_text)

        # Dodaj spację przed półpauzą
        matches11 = re.findall(r'(?<=\S)–', paragraph.text)
        for match in matches11:
            new_text = " –"
            paragraph.text = paragraph.text.replace(match, new_text)

        # Dodaj spację przed nawiasem
        matches12 = re.findall(r'(?<=[^\s])([\(\[])', paragraph.text)
        for match in matches12:
            new_text = " " + match
            paragraph.text = paragraph.text.replace(match, new_text, 1)

        # Usuń spację przed nawiasem zamykającym
        matches13 = re.findall(r'(\s[\)\]])', paragraph.text)
        for match in matches13:
            new_text = match.strip()
            paragraph.text = paragraph.text.replace(match, new_text, 1)

        # Dodaj kropkę na końcu akapitu, jeśli nie ma innego znaku niż litera albo cyfra
        if re.search(r'[\wąćęłńóśźżĄĆĘŁŃÓŚŹŻ0-9]$', paragraph.text, re.UNICODE):
            paragraph.text += '.'

        # Zamień amerykańskie cudzysłowy na polskie
        paragraph.text = re.sub(r'"(?=\w)', '„', paragraph.text)
        paragraph.text = re.sub(r'"(?=\W)', '”', paragraph.text)

        # Zamień polski cudzysłów zamykający na początku słowa na otwierający
        paragraph.text = re.sub(r'”(?=\w)', '„', paragraph.text, flags=re.IGNORECASE)

        # Zamień nieprawidłowy cudzysłów zamykający na początku słowa i zamień go na dolny otwierający
        paragraph.text = re.sub(r'“(?=\w)', '„', paragraph.text, flags=re.IGNORECASE)

        # Przenieś kropkę poza cudzysłów.
        paragraph.text = re.sub(r'\.”', '”.', paragraph.text)

        # Usuń wielokrotne spacje
        matches14 = re.findall(r'\s{2,}', paragraph.text)
        for match in matches14:
            new_text = ' '
            paragraph.text = paragraph.text.replace(match, new_text)

        # Zastąp małe litery na początku akapitów wielkimi
        matches15 = re.findall(r'^[a-ząćęłńóśźż]', paragraph.text)
        for match in matches15:
            new_text = match.upper()
            paragraph.text = paragraph.text.replace(match, new_text, 1)

        # Znajdź małe litery po pytajniku, wykrzykniku lub kropce
        matches16 = re.findall(r'(?<=[?.!])\s[a-ząćęłńóśźż]', paragraph.text)
        for match in matches16:
            # Zmień małą literę na dużą
            paragraph.text = re.sub(r'(?<=[?.!])\s[a-ząćęłńóśźż]', lambda x: x.group().upper(), paragraph.text, 1)

        # Usuń spacje na początku dokumentu
        if doc.paragraphs: 
            first_paragraph = doc.paragraphs[0]
            first_paragraph.text = first_paragraph.text.lstrip()
    
    base = os.path.basename(file_path) 
    name, ext = os.path.splitext(base) 
    new_name = f"{name}_poprawki{ext}" 

    # Zapisz z nową nazwą
    doc.save(new_name)

def main():
    window = tk.Tk()
    
    message = "Kontakt: dawidwiktorski@gmail.com\n\nWersja 0.1 (19.07.2023 r.)\n\nWYŁĄCZNIE pliki DOCX (nie mogą być otwarte w czasie wykonywania skryptu).\n\nAplikacja nie nadpisuje oryginalnego pliku i tworzy kopię w miejscu, w którym się znajduje (aplikacja).\n\nMimo to zalecam stworzyć kopię bezpieczeństwa.\n\nZe względu na charakter poprawek najlepiej wykonywać je na pliku przed czytaniem, w kodzie nie da się wszystkiego przewidzieć.\n\nUWAGA! Poprawione fragmenty stracą swoje formatowanie, należy je przywrócić ręcznie.\n\nGdy skrypt zakończy pracę, zostanie wyświetlony odpowiedni komunikat."

    msg_label = tk.Label(window, text=message)
    msg_label.pack()

    def on_click():
        file_path = open_file_dialog()
        if file_path:
            try:
                find_and_replace(file_path)
                messagebox.showinfo("Sukces", "Skrypt zakończył pracę")
            except Exception as e:
                messagebox.showerror("Błąd", f"Wystąpił błąd: {e}")

    open_button = tk.Button(window, text="Otwórz plik", command=on_click)
    open_button.pack()

    window.mainloop()

if __name__ == "__main__":
    main()