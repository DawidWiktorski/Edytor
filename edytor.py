import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
import os
import re

class Edytor:
    def __init__(self, root):
        self.root = root
        self.root.title("Edytor DOCX")
        self.root.geometry("800x800")
        
        # Lista polskich zaimków osobowych
        self.polish_pronouns = r'\b(Ja|Ty|On|Ona|Ono|My|Wy|Oni|One|Mnie|Ciebie|Jego|Jej|Nas|Was|Ich|Je|Mi|Ci|Mu|Nam|Wam|Im|Mną|Tobą|Nim|Nią|Nami|Wami|Nimi)\b'
        
        self.selected_file = None
        self.options = {
            # Spacje
            "Zamień wielokrotne spacje na pojedynczą": tk.BooleanVar(),
            "Usuń spacje z początku i końca akapitu": tk.BooleanVar(),
            "Usuń spacje przed znakami interpunkcyjnymi": tk.BooleanVar(),
            "Dodaj spację po znakach interpunkcyjnych (jeśli brakuje)": tk.BooleanVar(),
            "Dodaj spacje przed nawiasami i po nawiasach": tk.BooleanVar(),
            
            # Znaki interpunkcyjne
            "Zamień trzy kropki na wielokropek": tk.BooleanVar(),
            "Usuń kropkę po znaku wielokropka": tk.BooleanVar(),
            "Zamień zwielokrotnione znaki interpunkcyjne na pojedyncze": tk.BooleanVar(),
            "Dodaj kropkę na końcu akapitu (uwaga, doda też w nagłówkach!)": tk.BooleanVar(),
            "Przesuń kropkę za nawias": tk.BooleanVar(),
            
            # Dywizy i pauzy - półpauzy
            "Zamień podwójne dywizy na półpauzy": tk.BooleanVar(),
            "Zamień pauzy na półpauzy": tk.BooleanVar(),
            "Zamień dywizy na półpauzy (nie uwzględnia łączników)": tk.BooleanVar(),
            "Dodaj spacje do półpauz, jeśli brakuje": tk.BooleanVar(),            
            "Zamień dywiz otwierający akapit na półpauzę": tk.BooleanVar(),
            "Wielka litera po półpauzie": tk.BooleanVar(),
            
            # Cudzysłowy
            "Zamień amerykański cudzysłów na polski": tk.BooleanVar(),
            "Popraw cudzysłów otwierający i zamykający": tk.BooleanVar(),
            "Przenieś kropkę poza cudzysłów": tk.BooleanVar(),
                       
            # Formatowanie tekstu
            "Popraw wielkość liter w zaimkach osobowych": tk.BooleanVar(),
            "Wielka litera na początku zdania": tk.BooleanVar(),
        }
        
        self.create_widgets()

    def create_widgets(self):
        # Główny kontener
        main_container = ttk.Frame(self.root)
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Kontener dla lewej strony (główna zawartość)
        left_container = ttk.Frame(main_container)
        left_container.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Kontener dla prawej strony (ostrzeżenie)
        right_container = ttk.Frame(main_container, width=200)
        right_container.pack(side=tk.RIGHT, fill=tk.Y, padx=10)
        
        warning_frame = ttk.LabelFrame(right_container, text="Ostrzeżenie", padding="5")
        warning_frame.pack(fill=tk.X, pady=5)
        warning_label = ttk.Label(warning_frame, text="Program może wprowadzić\nbłędne poprawki.\nProszę o dokładne\nsprawdzenie zmian.\nProgram resetuje formatowanie tekstu!", 
                                wraplength=180, justify=tk.CENTER)
        warning_label.pack(pady=5)

        canvas = tk.Canvas(left_container)
        scrollbar = ttk.Scrollbar(left_container, orient="vertical", command=canvas.yview)
        
        self.scrollable_frame = ttk.Frame(canvas)
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        file_frame = ttk.Frame(self.scrollable_frame, padding="5")
        file_frame.pack(fill=tk.X, pady=5)
        
        self.file_label = ttk.Label(file_frame, text="Nie wybrano pliku")
        self.file_label.pack(side=tk.LEFT, padx=5)
        
        select_button = ttk.Button(file_frame, text="Wybierz plik DOCX", command=self.select_file)
        select_button.pack(side=tk.RIGHT)

        options_frame = ttk.LabelFrame(self.scrollable_frame, text="Opcje poprawek", padding="5")
        options_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # Checkboxy dla wszystkich opcji
        for option_text, var in self.options.items():
            cb = ttk.Checkbutton(options_frame, text=option_text, variable=var)
            cb.pack(anchor=tk.W, pady=2)

        # Przyciski kontrolne
        control_frame = ttk.Frame(self.scrollable_frame)
        control_frame.pack(fill=tk.X, pady=5)
        
        select_all_button = ttk.Button(control_frame, text="Zaznacz/odznacz wszystko", 
                                     command=self.toggle_all)
        select_all_button.pack(side=tk.LEFT, padx=5)
        
        process_button = ttk.Button(control_frame, text="Rozpocznij wykonywanie zmian", 
                                  command=self.process_file)
        process_button.pack(side=tk.RIGHT, padx=5)

        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.bind_mousewheel(canvas)

    def bind_mousewheel(self, widget):
        def _on_mousewheel(event):
            widget.yview_scroll(int(-1*(event.delta/120)), "units")
        widget.bind_all("<MouseWheel>", _on_mousewheel)

    def select_file(self):
        filename = filedialog.askopenfilename(
            filetypes=[("Dokumenty Word", "*.docx")],
            title="Wybierz plik DOCX"
        )
        if filename:
            self.selected_file = filename
            self.file_label.config(text=os.path.basename(filename))

    def toggle_all(self):
        # Sprawdź stan wszystkich opcji
        all_selected = all(var.get() for var in self.options.values())
        # Ustaw przeciwny stan dla wszystkich opcji
        new_state = not all_selected
        for var in self.options.values():
            var.set(new_state)

    def process_file(self):
        if not self.selected_file:
            messagebox.showerror("Błąd", "Najpierw wybierz plik!")
            return
            
        if not any(var.get() for var in self.options.values()):
            messagebox.showerror("Błąd", "Wybierz przynajmniej jedną opcję!")
            return
            
        try:
            doc = Document(self.selected_file)
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    new_text = self.process_text(paragraph.text)
                    paragraph.clear()
                    paragraph.add_run(new_text)
            
            file_path, file_extension = os.path.splitext(self.selected_file)
            new_file_path = f"{file_path}_edytowany{file_extension}"
            
            doc.save(new_file_path)
            messagebox.showinfo("Sukces", "Plik został pomyślnie zapisany!")
            
        except Exception as e:
            messagebox.showerror("Błąd", f"Wystąpił błąd: {str(e)}")

    def process_text(self, text):
        # Znaki interpunkcyjne
        if self.options["Zamień trzy kropki na wielokropek"].get():
            text = re.sub(r'\.{3}', '…', text)
            
        if self.options["Usuń kropkę po znaku wielokropka"].get():
            text = re.sub(r'…\.', '…', text)
            
        if self.options["Zamień zwielokrotnione znaki interpunkcyjne na pojedyncze"].get():
            text = re.sub(r'([.!?,:;])\1+', r'\1', text)

        if self.options["Przesuń kropkę za nawias"].get():
            text = re.sub(r'(\.)(\s*")\s*(?!\.)', r'\2\1', text)
        
        if self.options["Dodaj kropkę na końcu akapitu (uwaga, doda też w nagłówkach!)"].get():
            if text and not text.rstrip().endswith(('.', '!', '?', ':', ';', '…')):
                text = text.rstrip() + '.'

        # Półpauzy
        if self.options["Zamień podwójne dywizy na półpauzy"].get():
            text = text.replace('--', '–')
        
        if self.options["Zamień pauzy na półpauzy"].get():
            text = text.replace('—', '–')

        if self.options["Zamień dywizy na półpauzy (nie uwzględnia łączników)"].get():          
            text = re.sub(r'(?<![a-zA-Z0-9])-|-(?![a-zA-Z0-9])', '–', text)

        if self.options["Dodaj spacje do półpauz, jeśli brakuje"].get():
            # Zamień dywizy na półpauzy, z wyjątkiem tych między literami/cyframi
            text = re.sub(r'(?<![a-zA-Z0-9])-(?![a-zA-Z0-9])', '–', text)
            # Dodaj spacje przy półpauzach z zachowaniem wyjątków:
            text = re.sub(r'(\S)–(\S)', r'\1 – \2', text)  # Brak spacji z obu stron
            text = re.sub(r'(\s|^)–(\S)', r'\1 – \2', text)  # Brak spacji z prawej (bez początku akapitu)
            text = re.sub(r'(\S)–(\s|$)', r'\1 – \2', text)  # Brak spacji z lewej (bez końca akapitu)

        # Cudzysłowy
        if self.options["Zamień amerykański cudzysłów na polski"].get():
            text = re.sub(r'(\s|^)"', r'\1„', text)
            text = re.sub(r'^"', '„', text)
            text = re.sub(r'"([.,!?:;\s]|$)', r'”\1', text)

        if self.options["Popraw cudzysłów otwierający i zamykający"].get():
            text = re.sub(r'(\s)”', r'\1„', text)
            text = re.sub(r'([a-zA-Z0-9])„([.,!?:;\s]|$)', r'\1”\2', text)


        if self.options["Przenieś kropkę poza cudzysłów"].get():
            # Usuń kropkę sprzed cudzysłowu zamykającego
            text = re.sub(r'\.(”)', r'\1', text)
            
            # Dodaj kropkę po cudzysłowie zamykającym, jeśli następne słowo zaczyna się wielką literą
            text = re.sub(r'”(\s+)([A-ZĘÓĄŚŁŻŹĆŃ])', r'”\1.\2', text)

        # Nawiasy
        if self.options["Dodaj spacje przed nawiasami i po nawiasach"].get():
            text = re.sub(r'(\S)([\(\[\{])', r'\1 \2', text)
            text = re.sub(r'([\(\[\{])(\s+)', r'\1', text)
            text = re.sub(r'\s+([\)\]\}])', r'\1', text)
            text = re.sub(r'([\)\]\}])(\w)', r'\1 \2', text)

        # Spacje
        if self.options["Zamień wielokrotne spacje na pojedynczą"].get():
            text = re.sub(r' +', ' ', text)
            
        if self.options["Usuń spacje z początku i końca akapitu"].get():
            text = text.strip()
            
        if self.options["Usuń spacje przed znakami interpunkcyjnymi"].get():
            text = re.sub(r'\s+([.,!?:;])', r'\1', text)
            
        if self.options["Dodaj spację po znakach interpunkcyjnych (jeśli brakuje)"].get():
            text = re.sub(r'([.,!?:;])([a-zA-Z0-9])', r'\1 \2', text)
    
        # Wielkie litery
        if self.options["Wielka litera po półpauzie"].get():
            text = re.sub(r'^[–]\s*([a-ząćęłńóśźż])',
                        lambda m: m.group(0)[:-1] + m.group(1).upper(),
                        text, flags=re.MULTILINE)
            
        if self.options["Popraw wielkość liter w zaimkach osobowych"].get():
            def replace_pronoun(match):
                pronoun = match.group(0)
                start_pos = match.start()
                if start_pos == 0 or text[start_pos-2:start_pos].rstrip()[-1] in '.!?':
                    return pronoun
                return pronoun.lower()
            text = re.sub(self.polish_pronouns, replace_pronoun, text)

        if self.options["Wielka litera na początku zdania"].get():
            if text and text[0].isalpha():
                text = text[0].upper() + text[1:]
            
            # Oznaczenie wystąpień
            text = re.sub(
                r'(\b(?:np|tzw)\.\s+)([a-ząćęłńóśźż])',
                r'\1#KEEP#\2',  # Dodanie znacznika
                text,
                flags=re.IGNORECASE
            )
            
            text = re.sub(
                r'([.!?]\s+)(?!#KEEP#)([a-ząćęłńóśźż])',
                lambda m: m.group(1) + m.group(2).upper(),
                text
            )
            
            text = text.replace('#KEEP#', '')

        return text

def main():
    root = tk.Tk()
    app = Edytor(root)
    root.mainloop()

if __name__ == "__main__":
    main()